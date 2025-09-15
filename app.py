import streamlit as st
import pdfplumber
from docx import Document
from fpdf import FPDF
from datetime import datetime
import os

st.set_page_config(page_title="🏛 Bank Legal Opinion Tool", layout="wide")
st.title("🏛 Bank Legal Opinion Tool")
st.write("Upload property documents and generate legal opinion (Word & PDF).")

uploaded_files = st.file_uploader("📄 Upload property PDF documents", type="pdf", accept_multiple_files=True)
generate_btn = st.button("Generate Legal Opinion")

if generate_btn:
    if not uploaded_files:
        st.error("Please upload at least one PDF document.")
    else:
        extracted_text = ""
        for file in uploaded_files:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"

        # Load template
        template = Document("template.docx")
        today = datetime.today().strftime("%d-%m-%Y")

        # Replace placeholders in Word template
        for p in template.paragraphs:
            if "{{DATE}}" in p.text:
                p.text = p.text.replace("{{DATE}}", today)
            if "{{PROPERTY_DETAILS}}" in p.text:
                p.text = p.text.replace("{{PROPERTY_DETAILS}}", extracted_text)

        # Save Word file
        filled_docx = "filled_opinion.docx"
        template.save(filled_docx)

        st.success("✅ Word document generated!")
        with open(filled_docx, "rb") as f:
            st.download_button("⬇️ Download Word File", f, file_name="Legal_Opinion.docx")

        # Generate PDF from Word content
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Times", size=12)

        for paragraph in template.paragraphs:
            pdf.multi_cell(0, 10, paragraph.text)

        output_pdf = "final_opinion.pdf"
        pdf.output(output_pdf)

        st.success("✅ PDF generated!")
        with open(output_pdf, "rb") as f:
            st.download_button("⬇️ Download PDF File", f, file_name="Legal_Opinion.pdf")